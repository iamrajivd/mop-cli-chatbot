
import os
import streamlit as st
from datetime import datetime
from langchain.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain.chains import RetrievalQA
from langchain_community.chat_models import ChatOllama
from docx import Document
from fpdf import FPDF

# Configuration
EMBED_MODEL = "all-minilm"
VECTOR_DIRS = {
    "CMM": "vector_store/cmm_index",
    "CMG": "vector_store/cmg_index",
}
os.makedirs("vector_store", exist_ok=True)

# Load FAISS index
def load_index(domain):
    emb = OllamaEmbeddings(model=EMBED_MODEL)
    return FAISS.load_local(VECTOR_DIRS[domain], emb, allow_dangerous_deserialization=True)

# Determine query domain
def detect_domain(query):
    q = query.lower()
    if any(w in q for w in ["cmg", "pgw", "gtp", "sgw", "gprs"]):
        return "CMG"
    return "CMM"

# Build fast RAG agent
def get_agent(domain):
    vector = load_index(domain)
    retriever = vector.as_retriever(search_type="similarity", k=3)
    llm = ChatOllama(model="llama3", temperature=0.2)
    return RetrievalQA.from_chain_type(llm=llm, retriever=retriever, return_source_documents=False)

# Streamlit UI
st.set_page_config("RAG CLI Chatbot", layout="wide")
st.title("⚡ Fast CLI Chatbot (RAG + LLaMA3)")

with st.sidebar:
    st.header("📋 MoP Metadata")
    engineer = st.text_input("Engineer Name")
    cr_id = st.text_input("CR Number")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

if "chat" not in st.session_state:
    st.session_state.chat = []

query = st.chat_input("Ask CLI-specific query (e.g. patchParameter or gParms)...")
if query:
    domain = detect_domain(query)
    agent = get_agent(domain)

    st.session_state.chat.append(("user", f"[{domain} DOMAIN]\n{query}"))

    with st.spinner(f"🔍 Searching {domain} CLI..."):
        response = agent.run({"query": query})
    st.session_state.chat.append(("agent", response))

    with open("query_log.txt", "a") as f:
        f.write(f"[{timestamp}] {engineer}/{cr_id} - {domain}\nQ: {query}\nA: {response}\n{'='*60}\n")

for role, msg in st.session_state.chat:
    st.chat_message(role).write(msg)

if st.button("📁 Export MoP"):
    mop_text = f"Engineer: {engineer}\nCR Number: {cr_id}\nTimestamp: {timestamp}\n\n"
    mop_text += "\n\n".join([f"{role.upper()}: {msg}" for role, msg in st.session_state.chat])

    doc = Document()
    doc.add_heading("Method of Procedure (MoP)", 0)
    doc.add_paragraph(mop_text)
    doc.save("MoP_RAG_CLI.docx")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in mop_text.splitlines():
        pdf.cell(200, 10, txt=line, ln=True)
    pdf.output("MoP_RAG_CLI.pdf")

    st.download_button("⬇ Word", data=open("MoP_RAG_CLI.docx", "rb"), file_name="MoP_RAG_CLI.docx")
    st.download_button("⬇ PDF", data=open("MoP_RAG_CLI.pdf", "rb"), file_name="MoP_RAG_CLI.pdf")
